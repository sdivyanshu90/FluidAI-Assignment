from __future__ import annotations

import asyncio
from pathlib import Path
import pytest
from docx import Document
from app.agent import DocumentAgent
from app.config import Settings
from tests.fakes import FakeGeminiProvider

STANDARD_REQUEST = (
    "Create a project plan for launching an employee wellness program for 500 staff "
    "within 12 weeks with scope, owners, risks, and success metrics."
)
COMPLEX_REQUEST = (
    "Prepare a technical design and rollout plan for AI-assisted customer onboarding "
    "in 8 weeks with missing compliance requirements and conflicting rollout constraints."
)

def run_agent(tmp_path: Path, input_text: str):
    settings = Settings(gemini_api_key="test-key", output_dir=tmp_path)
    return asyncio.run(
        DocumentAgent(settings, FakeGeminiProvider()).run(input_text)
    )

@pytest.mark.parametrize("input_text", [STANDARD_REQUEST, COMPLEX_REQUEST])
def test_agent_generates_readable_docx_with_trace(
    tmp_path: Path, input_text: str
) -> None:
    response = run_agent(tmp_path, input_text)
    artifact = tmp_path / response.document_name

    assert artifact.is_file()
    assert artifact.stat().st_size > 10_000
    assert response.provider == "gemini/test-double"
    assert response.plan[-1].tool.value == "render_document"
    assert all(task.status == "completed" for task in response.plan)
    assert len(response.guardrails_applied) >= 6

    document = Document(artifact)
    content = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Executive Summary" in content
    assert "Assumptions" in content
    assert len(document.tables) >= 4

def test_complex_request_adds_assumption_task(tmp_path: Path) -> None:
    response = run_agent(tmp_path, COMPLEX_REQUEST)

    assert any(task.tool.value == "record_assumptions" for task in response.plan)
    assert len(response.assumptions) == 2

def test_standard_request_uses_minimal_safe_plan(tmp_path: Path) -> None:
    response = run_agent(tmp_path, STANDARD_REQUEST)

    assert [task.tool.value for task in response.plan] == [
        "analyze_request",
        "draft_document",
        "render_document",
    ]
