from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from app.schemas import DocumentDraft

NAVY = "17365D"
BLUE = "2F75B5"
PALE_BLUE = "D9EAF7"
LIGHT_GREY = "F2F4F7"
WHITE = "FFFFFF"
DARK_GREY = RGBColor(55, 65, 81)

def safe_document_name(title: str, job_id: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", title).strip("-").lower()[:55]
    return f"{slug or 'document'}-{job_id[:8]}.docx"

def render_docx(draft: DocumentDraft, destination: Path) -> None:
    """Render a business-ready Word document from the agent's structured draft."""

    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document)
    _add_cover(document, draft)
    _add_contents(document, draft)

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(draft.executive_summary, style="Executive Summary")

    for section in draft.sections:
        document.add_heading(section.heading, level=1)
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)
        for bullet in section.bullets:
            document.add_paragraph(bullet, style="List Bullet")
        if section.table:
            if section.table.title:
                caption = document.add_paragraph(section.table.title)
                caption.style = "Caption"
            _add_table(document, section.table.headers, section.table.rows)

    if draft.assumptions:
        document.add_heading("Assumptions", level=1)
        document.add_paragraph(
            "These statements enable planning and require stakeholder validation."
        )
        for assumption in draft.assumptions:
            document.add_paragraph(assumption, style="List Bullet")

    _add_signoff(document)
    document.core_properties.title = draft.title
    document.core_properties.subject = draft.subtitle or "AI-generated business document"
    document.core_properties.author = "Autonomous Document Agent"
    document.core_properties.keywords = "autonomous agent, business document"
    document.save(destination)

def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK_GREY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for level, size, color in ((1, 18, NAVY), (2, 14, BLUE), (3, 11, NAVY)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(15 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)

    if "Executive Summary" not in document.styles:
        summary_style = document.styles.add_style(
            "Executive Summary", WD_STYLE_TYPE.PARAGRAPH
        )
    else:
        summary_style = document.styles["Executive Summary"]
    summary_style.base_style = normal
    summary_style.font.size = Pt(11)
    summary_style.font.italic = True
    summary_style.font.color.rgb = RGBColor.from_string(NAVY)
    summary_style.paragraph_format.left_indent = Inches(0.25)
    summary_style.paragraph_format.right_indent = Inches(0.25)
    summary_style.paragraph_format.space_after = Pt(12)

    _set_header_footer(section)

def _add_cover(document: Document, draft: DocumentDraft) -> None:
    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.cell(0, 0)
    _shade_cell(cell, NAVY)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    label = cell.paragraphs[0]
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("AUTONOMOUS DOCUMENT AGENT")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(WHITE)

    document.add_paragraph("\n\n")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(draft.title)
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    if draft.subtitle:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(draft.subtitle)
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor.from_string(BLUE)

    document.add_paragraph("\n")
    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_run = rule.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    rule_run.font.color.rgb = RGBColor.from_string(BLUE)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run("Prepared for\n").bold = True
    metadata.add_run(f"{draft.audience}\n\n")
    metadata.add_run("Prepared on\n").bold = True
    metadata.add_run(datetime.now(timezone.utc).strftime("%d %B %Y"))

    document.add_page_break()

def _add_contents(document: Document, draft: DocumentDraft) -> None:
    document.add_heading("Document Guide", level=1)
    document.add_paragraph(
        "This guide provides a quick view of the document's decision structure."
    )
    entries = ["Executive Summary", *(section.heading for section in draft.sections)]
    if draft.assumptions:
        entries.append("Assumptions")
    for index, heading in enumerate(entries, start=1):
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(heading)
    document.add_page_break()

def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            if row_index % 2:
                _shade_cell(cells[column_index], LIGHT_GREY)
            paragraph = cells[column_index].paragraphs[0]
            run = paragraph.add_run(value)
            run.font.size = Pt(8.5)
    document.add_paragraph()

def _add_signoff(document: Document) -> None:
    document.add_heading("Approval", level=1)
    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ["Role", "Name / Signature", "Date"]
    for index, value in enumerate(headers):
        _shade_cell(table.cell(0, index), PALE_BLUE)
        table.cell(0, index).paragraphs[0].add_run(value).bold = True
    table.cell(1, 0).text = "Document owner"
    table.cell(1, 1).text = ""
    table.cell(1, 2).text = ""

def _set_header_footer(section: WD_SECTION) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("AUTONOMOUS DOCUMENT AGENT  •  WORKING DOCUMENT")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = footer.add_run("Generated document  |  Page ")
    label.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

def _shade_cell(cell: object, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)

